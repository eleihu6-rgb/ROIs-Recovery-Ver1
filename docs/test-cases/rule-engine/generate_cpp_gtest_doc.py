"""Generate the C++ gtest validation-logic summary as a .docx (Ver 2).

Reproducible source for `cpp-gtest-validation-logic-Ver2.docx`. Run inside the
rule-engine venv (which has python-docx):

    cd rule-engine && source .venv/bin/activate
    python ../docs/test-cases/rule-engine/generate_cpp_gtest_doc.py

Ver 2 adds: a "Test methodology" section (specification-based / equivalence
partitioning, standard-term table) and, for every rule, a plain-language
positive (legal) and negative (illegal) case explanation alongside the tech specs.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

OUT = Path(__file__).with_name("cpp-gtest-validation-logic-Ver2.docx")

GREY = RGBColor(0x59, 0x59, 0x59)
GREEN = RGBColor(0x1E, 0x7A, 0x34)
RED = RGBColor(0xB0, 0x2A, 0x2A)


def h(doc, text, level):
    return doc.add_heading(text, level=level)


def kv_table(doc, rows):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v
        for para in cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for para in cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()
    return t


def case_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light List Accent 1"
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = hd
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()
    return t


def plain_language(doc, pos, neg):
    """Render the user-friendly positive (legal) / negative (illegal) story."""
    p = doc.add_paragraph()
    r = p.add_run("✅ Positive case (legal). ")
    r.font.bold = True
    r.font.color.rgb = GREEN
    p.add_run(pos)
    p = doc.add_paragraph()
    r = p.add_run("❌ Negative case (illegal). ")
    r.font.bold = True
    r.font.color.rgb = RED
    p.add_run(neg)


doc = Document()

# ── Title ───────────────────────────────────────────────────────────────────
doc.add_heading("ROIS Rule Engine — C++ gtest Validation Logic (Ver 2)", level=0)
sub = doc.add_paragraph(
    "How the legacy C++ rules (crewrule-dev/RuleTest/) validate each F8 regulation, "
    "and how those scenarios were replicated to validate the Python migration."
)
for run in sub.runs:
    run.font.italic = True
    run.font.color.rgb = GREY
doc.add_paragraph("Generated 2026-06-14 · 6 of the 14 F8 rules have a behavioral C++ gtest.")

# ── 1. Test methodology ─────────────────────────────────────────────────────
h(doc, "1. Test methodology", 1)
doc.add_paragraph(
    "Each rule is tested with specification-based (black-box) tests using equivalence "
    "partitioning: the input domain is split into a legal partition and an illegal partition. "
    "A positive test case supplies a legal roster and asserts the legal verdict (no violation); "
    "a negative test case supplies an illegal roster and asserts the illegal verdict (violation "
    "raised). The expected legal/illegal outcome is the test oracle, and the legal/illegal flag "
    "is the pass/fail criterion."
)
doc.add_paragraph(
    "Two supporting techniques recur: Boundary Value Analysis — testing right at the threshold "
    "(e.g. 6715 min just-legal vs 6725 min just-illegal around the 6720-min / 112:00 limit) — and "
    "two-sided decision coverage: pairing a positive and a negative case exercises both outcomes of "
    "the rule's predicate, which is what prevents a degenerate “always-legal” or "
    "“always-illegal” implementation from passing."
)
h(doc, "Plain wording → standard term", 3)
case_table(
    doc,
    ["Plain wording", "Standard test term"],
    [
        ["Green case, expect no violation, mark roster legal",
         "Positive test case — valid input / legal partition; oracle = pass"],
        ["Red case, expect violation, mark roster illegal",
         "Negative test case — invalid input / illegal partition; oracle = fail"],
        ["Splitting rosters into legal vs illegal", "Equivalence partitioning"],
        ["Testing right at the limit (6715 vs 6725)", "Boundary value analysis (BVA)"],
        ["Having both a green and a red case", "Two-sided / decision coverage"],
        ["The expected legal/illegal answer", "Test oracle"],
        ["The legal/illegal flag that decides pass/fail", "Pass/fail criterion"],
    ],
)
doc.add_page_break()

# ── 2. Common harness ───────────────────────────────────────────────────────
h(doc, "2. The common gtest harness (shared by every rule)", 1)
doc.add_paragraph(
    "Each rule*_gtest.cpp follows the same five-step shape. Understanding it once "
    "explains every per-rule test below."
)
steps = [
    ("buildContext(crewId)", "Creates a CrewDataContext: airline (F8), scenario start/end dates, "
        "base airport(s) with their UTC offset and IANA zone (e.g. YEG = UTC-360 / America/Edmonton), "
        "and a CREW with a CREW_BASE. This is the world the rule runs in."),
    ("makeXxxRoster(...)", "Builds ROSTER rows (the crew's assignments/duties): start/end/rest times in "
        "BOTH UTC and crew-base local, a qualifier (GND/GRD/PRPM…) and a duty group (FLY/DO…). "
        "The local times are what most time-window rules actually evaluate."),
    ("makeXxxxInput(...)", "Builds a RuleInput holding a DBRule with the regulation PARAMETERS in a "
        "params map (e.g. MIN PERIOD=55, UNIT=RH, MIN DO=12). This is the rule's configuration."),
    ("rule->CheckRule(rosters)", "Runs the rule over the roster list. Returns bool and fills a "
        "_violations vector. (Calculator rules instead call rule->Calculate(...) and return a number.)"),
    ("EXPECT_TRUE / EXPECT_FALSE", "PASS (legal): CheckRule()==true AND _violations empty. "
        "FAIL (illegal): CheckRule()==false AND _violations non-empty (often with specific fields asserted)."),
]
for name, desc in steps:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(name + " — ")
    r.font.bold = True
    r.font.name = "Consolas"
    p.add_run(desc)
note = doc.add_paragraph()
nr = note.add_run("In one line: a positive (green) test feeds the rule a LEGAL roster and asserts no "
                  "violation; a negative (red) test feeds an ILLEGAL roster and asserts the violation fires. "
                  "Each rule ships both.")
nr.font.italic = True
doc.add_page_break()

# ── Per-rule sections ───────────────────────────────────────────────────────
RULES = [
    {
        "title": "3. Rule 7506/002 — Single Daily Check-in",
        "meta": [
            ("C++ file / class", "rule7506_gtest.cpp · SingleDailyCheckinForCARSRule"),
            ("Python target", "single_daily_checkin_checker.py · SingleDailyCheckinChecker"),
            ("Rule intent", "Only one duty check-in (roster START) is allowed per LOCAL calendar day."),
            ("Migration verdict", "FAITHFUL (4/4 replicated). ASSIGNMENTS filter unused by F8 (empty params)."),
        ],
        "data": "World: airline F8, single base AAA at UTC offset 0 (so local day == UTC day). One crew "
                "with base AAA. Inputs: ground rosters (qualifier GND) at AAA, each with a start/end/rest "
                "time. Rule params: ASSIGNMENTS = checked groups (default 'FLY'); BASES/RANKS/FLEETS = '*'.",
        "how": "CheckRule scans the roster starts and groups them by local calendar day. PASS when no two "
               "CHECKED check-ins share a day; FAIL (violation) when two do. The rule keys on the START "
               "time only, not the end.",
        "pos": "A crew member checks in once on Monday and once on Tuesday — one check-in each day. "
               "Nothing is wrong, so the roster is marked legal and no warning appears.",
        "neg": "The same crew member is rostered to start two separate duties on the SAME Monday "
               "(a morning duty and an evening duty). You can't legally begin two duties in one calendar "
               "day, so the rule marks the roster illegal and raises a violation.",
        "cases_h": ["Test case", "Input data", "Expected"],
        "cases": [
            ["AllowsDifferentStartLocalDays", "Starts Mar-01 08:00 and Mar-02 08:00", "PASS (no violation)"],
            ["RejectsSameStartLocalDay", "Starts Mar-01 08:00 and Mar-01 20:00", "FAIL (violation)"],
            ["AllowsSameEndAndNextStartDayWhenStartsDiffer",
             "Duty1 Mar-01 08:00 → Mar-02 06:00; Duty2 Mar-02 07:00 (ends/starts share Mar-02, STARTS differ)",
             "PASS (rule keys on start)"],
            ["IgnoresUncheckedAssignmentGroups",
             "Two same-day duties, ASSIGNMENTS='DO', only one is DO (other FLY)", "PASS (only 1 checked)"],
        ],
    },
    {
        "title": "4. Rule 7504/003 — WOCL Spacing",
        "meta": [
            ("C++ file / class", "rule7504_gtest.cpp · CheckMinSpaceBetweenDutyForF8Rule"),
            ("Python target", "wocl_spacing_checker.py · WoclSpacingChecker"),
            ("Rule intent", "Min 55 rolling-hours between two duties that both touch the WOCL window (02:00–05:59)."),
            ("Migration verdict", "FIXED 2026-06-14 — WOCL membership now OVERLAP-based (was report-time-based)."),
        ],
        "data": "World: F8, base YEG at UTC-360 (MDT, America/Edmonton), May 2026, division P. Inputs: "
                "pre-assigned GRD reserve rosters (qualifier GRD, source PA) at YEG, given in both UTC and "
                "local. Rule params (7504003): LEVEL=D, MIN PERIOD=55, UNIT=RH (rolling hours), "
                "PREV/NEXT ATTRIBUTES=WOCL, APPLY PRELABELLED ATTRIBUTES=N, UTILIZE POST REST=Y.",
        "how": "A duty is 'WOCL' if its body OVERLAPS 02:00–05:59 in crew-base LOCAL time (the key fix: the "
               "old code compared raw UTC and mis-classified duties). If two WOCL duties are < 55 rolling-hours "
               "apart → violation. PASS when daytime duties are correctly NOT WOCL; FAIL when overnight WOCL "
               "duties are too close.",
        "pos": "Two afternoon reserve duties (2 pm to just-before-midnight) on back-to-back days. Neither "
               "overlaps the body-clock-low window (2–6 am), so the 55-hour spacing rule doesn't apply — the "
               "roster is legal and no warning shows.",
        "neg": "Two overnight reserves (midnight to 5 am) only ~19 hours apart. Both fall inside the "
               "window-of-circadian-low and they are far closer than the required 55 hours, so the rule marks "
               "the roster illegal and warns the duties are too tightly spaced.",
        "cases_h": ["Test case", "Input data", "Expected"],
        "cases": [
            ["DaytimeGroundReserveIsNotWoclAndPassesSpacing",
             "Two GRD at YEG local 14:00–23:59 (UTC 20:00 → 05:59 next day) on consecutive days",
             "PASS — 14:00–23:59 local is NOT WOCL (regression for the local-vs-UTC bug)"],
            ["OvernightGroundReserveIsWoclAndFailsSpacing",
             "Two GRD at YEG local 00:00–05:00 (overlap 02:00–05:59), gap ~19h",
             "FAIL — both WOCL, 19h < 55RH"],
        ],
    },
    {
        "title": "5. Rule 7505/002 — Min # GDOs in a Roster Period",
        "meta": [
            ("C++ file / class", "rule7505_gtest.cpp · MinimumDaysOffForCARSRule"),
            ("Python target", "min_gdo_checker.py · MinGdoChecker (legacy min_gdo / bid-period form)"),
            ("Rule intent", "At least 12 General Days Off (DO) per roster period (≈ one calendar month)."),
            ("Migration verdict", "FAITHFUL — recovered value 12 made explicit in 2026-06-14 migration."),
        ],
        "data": "World: F8, base YEG. Inputs: crew '247' with one DO on Jun-01 plus a daily PRPM standby "
                "(SBY) Jun 01–20 (pairing #32, YEG local 14:00). Rule params (7505002): DO ASSIGNMENT GROUP=DO, "
                "MIN DO=12, PERIOD=1, UNIT=RP, RP DAYS RANGE=30-30, COUNT BLANK DAY=Y, COUNT LAYOVER=N, "
                "LEAVE ASSIGNMENTS=VAC, UTILIZE POST DUTY REST=Y.",
        "how": "CheckRule counts DO days (plus blank days) within the 30-day roster period; standby (SBY) "
               "does NOT count as a day off. PASS when ≥12; FAIL when < 12, emitting a violation carrying "
               "rMinDO=12, rUnit=RP and iDaysOff = the actual (insufficient) count.",
        "pos": "Over the month the crew member gets 12 or more real days off (DO). The minimum is met, so the "
               "roster is legal with no warning.",
        "neg": "The crew member is put on standby almost every day in June and ends the month with fewer than "
               "12 actual days off. They are owed at least 12 per roster period, so the rule marks the roster "
               "illegal and reports how many days off were actually given (a number below 12).",
        "cases_h": ["Test case", "Input data", "Expected"],
        "cases": [
            ["RejectsCrew247DailyPrpmSbyJun2026InsufficientDaysOff",
             "20 days of SBY + a few late-June blank days → DO count < 12",
             "FAIL — violation with rMinDO=12, rUnit=RP, 0 ≤ iDaysOff < 12"],
        ],
    },
    {
        "title": "6. Rule 7272/001 — Calculate DP of the Reserves (Calculator)",
        "meta": [
            ("C++ file / class", "rule7272_gtest.cpp · CalculateStandbyDPForTGRule"),
            ("Python target", "duty_time_calculator.py · DutyTimeCalculator"),
            ("Rule intent", "Compute the duty-period (DP) credited for a reserve/standby block."),
            ("Migration verdict", "MIGRATED 2026-06-14 (was a sum-minutes stub). Callout add-on still out of scope."),
        ],
        "data": "World: YEG (UTC-360). Inputs: a single-duty standby pairing whose Duty/Segment carry an "
                "ASSIGNMENT qualifier (PRAM / SBY / RES). Rule params: ASSIGNMENTS (e.g. 'SBY|PRAM|PRPM'), "
                "RATE (0.33), STANDBY OFFSET ('HH:MM').",
        "how": "This is a CALCULATOR, not a checker: tests call rule->Calculate(roster, nextRoster) and assert "
               "the returned long = DP in SECONDS. Formula: DP_sec = floor( max(0, duration − offset) × rate ); "
               "returns −1 when the qualifier is NOT in ASSIGNMENTS (rule not applicable). No legal/illegal "
               "verdict — the assertion is on the computed number.",
        "pos": "A crew member sits a 12-hour standby of a type the rule covers (e.g. PRAM). The calculator "
               "works out how much of that standby counts as duty — about a third of it — and returns that "
               "credited time. (The “legal” equivalent here is “rule applies and computes a "
               "value”.)",
        "neg": "The standby is a type the rule does NOT cover (e.g. RES, not in the configured list). The "
               "calculator says “not my job” and returns −1, leaving the duty time untouched. (The "
               "“illegal/abnormal” equivalent here is “not applicable”, deliberately "
               "distinguished from a real zero.)",
        "cases_h": ["Test case", "Input / params", "Expected (Calculate return)"],
        "cases": [
            ["Calculate_Pram12hRate033_Returns14256Seconds", "PRAM 12h, RATE 0.33", "14256 s (= 43200×0.33)"],
            ["Calculate_SbyQualifier_MatchesProductionAssignmentList", "SBY 12h, in ASSIGNMENTS", "14256 s"],
            ["Calculate_UnlistedQualifier_ReturnsMinusOne", "RES 12h, not in list", "−1 (N/A)"],
            ["Calculate_AssignmentsSbyOnly_DoesNotMatchPramQualifier", "PRAM, ASSIGNMENTS='SBY'", "−1 (N/A)"],
            ["Calculate_WithOneHourOffset_ReducesDp", "PRAM 12h, STANDBY OFFSET 01:00", "13068 s"],
            ["Calculate_DurationBelowOffset_ReturnsZero", "PRAM 12h, OFFSET 13:00 > 12h", "0"],
            ["CalculatePairingDP_SetsDutyActualDpMinutes", "PRAM 12h pairing path", "ActualDP 237 min, DPInSecs 14220"],
            ["Calculate_Callout_AddsNotifyPeriodWithinLimit", "Standby + next FLY within notify limit",
             "7128 s (adds callout) — cross-pairing, xfail in Python"],
        ],
    },
    {
        "title": "7. Rule 8002/006 — Maximum Flight Time (cumulative)",
        "meta": [
            ("C++ file / class", "rule8002_gtest.cpp · (full crew legality, rule 8002 BH check)"),
            ("Python target", "rolling_flight_time_checker.py · MaxFtChecker (template max_ft)"),
            ("Rule intent", "Cumulative block hours must not exceed 112:00 in any rolling 28 calendar-day window."),
            ("Migration verdict", "FAITHFUL — 6715 passes / 6725 fails at the 6720-min (112h) limit."),
        ],
        "data": "World: 28-calendar-day window. Inputs: 15 realistic fly pairings (regional, red-eye, "
                "long-haul, multi-sector, ≥4 cross-midnight segments, one non-operating DHD sector whose "
                "block still counts). Block minutes tuned to a known total. Rule param: strMax='112:00' "
                "(6720 min), rolling 28 CD.",
        "how": "Full crew legality runs and the test looks for an 8002 'block hours' violation. PASS = legal "
               "and no 8002 BH violation. FAIL = not legal and a BH violation present (message contains "
               "'block hours'). The fixture also asserts the computed block total equals the intended value.",
        "pos": "Over a rolling 28 days a pilot flies 111 hours 55 minutes — just five minutes under the "
               "112-hour cap. They're within limits, so the roster is legal and no warning fires.",
        "neg": "The same 28 days but 112 hours 5 minutes of flying — five minutes over the cap. The rule marks "
               "the roster illegal and warns the pilot has exceeded the maximum block hours for the month. "
               "(This pair sits right on the boundary — classic boundary value analysis.)",
        "cases_h": ["Test case", "Input data", "Expected"],
        "cases": [
            ["TwentyEightDayRollingBh111h55mWithin112hLimitPasses",
             "Total block 6715 min (111:55), just under 6720", "PASS — no BH violation"],
            ["TwentyEightDayRollingBh112h05mExceeds112hLimitFails",
             "Total block 6725 min (112:05), 5 min over", "FAIL — BH violation ('block hours')"],
        ],
    },
    {
        "title": "8. Rule 7501/004 — Single Day Free From Duty (rolling windows)",
        "meta": [
            ("C++ file / class", "rule7501_gtest.cpp · LimitSingleDayFreeFromDutyForCARSRule"),
            ("Python target", "min_gdo_checker.py · MinGdoChecker (windows form)"),
            ("Rule intent", "≥1 SDFD in any rolling 168h window AND ≥4 SDFD in any rolling 672h window."),
            ("Migration verdict", "FAITHFUL for window evaluation; Python takes pre-computed DO days as input."),
        ],
        "data": "World: F8, base YEG. Inputs: rosters of duties separated by rest blocks; the C++ rule DERIVES "
                "an SDFD (Single Day Free from Duty) from rest that spans a local night. Rule params: two "
                "windows — PERIOD=168 UNIT=RH MIN LIMITS=1, and PERIOD=672 UNIT=RH MIN LIMITS=4, DUTY END "
                "BUFFER=00:30.",
        "how": "CheckRule slides each rolling window across the period and counts SDFDs; too few in any "
               "position → violation. PASS when every window meets its minimum; FAIL when a window is short. "
               "(Several C++ tests also cover optimizer vs editor source-masking and local-night derivation — "
               "C++-internal mechanics with no Python surface.)",
        "pos": "The crew gets a proper long rest each week, so there is always at least one full day free of "
               "duty in any 7-day stretch and at least four in any 28-day stretch. The roster is legal.",
        "neg": "The crew works eight duty days in a row with only short overnight rests — no full clear day off "
               "in the week. Everyone needs at least one free day per 168 hours, so the rule marks the roster "
               "illegal and warns a required day off is missing.",
        "cases_h": ["Test case", "Input data", "Expected"],
        "cases": [
            ["AllowsOneSdfdIn168Hours", "Long rests → ≥1 SDFD per 168h", "PASS"],
            ["RejectsMissingSdfdIn168Hours", "8 consecutive daily duties (~14h rest), no SDFD in a 168h window", "FAIL"],
            ["AllowsFourSdfdIn672Hours", "Weekly long rests → ≥4 SDFD per 672h", "PASS (both windows)"],
            ["June2026FiveOvernightDuties168Fails672Passes", "5 weekly-overnight blocks",
             "168h/1 FAILS, 672h/4 PASSES"],
        ],
    },
]

for rule in RULES:
    h(doc, rule["title"], 1)
    kv_table(doc, rule["meta"])
    p = doc.add_paragraph()
    p.add_run("1. Data used to validate the rule. ").bold = True
    p.add_run(rule["data"])
    p = doc.add_paragraph()
    p.add_run("2. How it is tested + pass/fail. ").bold = True
    p.add_run(rule["how"])
    h(doc, "In plain language", 3)
    plain_language(doc, rule["pos"], rule["neg"])
    h(doc, "Test cases", 3)
    case_table(doc, rule["cases_h"], rule["cases"])
    doc.add_page_break()

# ── Closing summary ─────────────────────────────────────────────────────────
h(doc, "9. Summary — migration fidelity (all 6 gtest-backed rules)", 1)
case_table(
    doc,
    ["Rule", "C++ gtest", "Python replica", "Verdict"],
    [
        ["7506 single_daily_checkin", "4 cases", "4 passed", "FAITHFUL"],
        ["7504 wocl_spacing", "2 cases", "fixed → passing", "FIXED (overlap)"],
        ["7505 min_gdo (RP)", "1 case", "3 passed", "FAITHFUL (value 12 recovered)"],
        ["7272 duty_period_calculator", "8 cases", "10 passed, 1 xfail", "MIGRATED (callout out of scope)"],
        ["8002 max_ft", "2 cases", "3 passed", "FAITHFUL"],
        ["7501 min_gdo (rolling)", "4 core cases", "4 passed", "FAITHFUL (window eval)"],
    ],
)
doc.add_paragraph(
    "The other 8 of the 14 F8 rules have no behavioral C++ gtest (only CSV params or SIA integration "
    "tests); acc_state, credit_hours_calc and roster_spacing were instead covered with spec-derived "
    "Python tests."
)

doc.save(str(OUT))
print(f"Wrote {OUT}")
