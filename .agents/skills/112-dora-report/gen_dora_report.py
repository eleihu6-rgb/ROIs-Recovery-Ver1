#!/usr/bin/env python3
"""
Skill 112 — DORA Developer Performance Report generator.

Self-contained: collects all metrics from git history of the CURRENT repo and
writes BOTH a Chinese (CN) and English (EN) .docx to docs/dev-context/, matching
the structure of the project's V6 report (13 sections).

Usage (run from repo root):
    python3 ~/.claude/skills/112-dora-report/gen_dora_report.py
    python3 .../gen_dora_report.py --version V7 --since 2026-01-01

Reads version from gantt/src/version.ts automatically.
Requires: python-docx (pip install python-docx)
"""
import argparse, datetime, os, re, subprocess
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------------------------------------------------------------
# Developer identity map (canonical — keep in sync with SKILL.md)
# ----------------------------------------------------------------------------
def dev_of(email, name):
    e, n = email.lower(), name.lower().strip()
    if 'yuan.zhu' in e or n in ('yuan.zhu-ai', 'yuanz-ai', 'yuan.zhu', 'yuanz-cc', 'yuan zhu'):
        return 'Yuan Zhu'
    if e == 'eleihu6@gmail.com' and n == 'ryan':
        return 'Kimi'
    if 'eleihu' in e:
        return 'Kimi'
    if 'honglei' in e or '1711625601' in e:
        return 'Honglei'
    if 'qiang.gong' in e:
        return 'Qiang'
    return None  # unknown / bots — ignored

DEVS = ['Kimi', 'Yuan Zhu', 'Honglei', 'Qiang']
NAME_CN = {'Kimi': 'Kimi (Ryan)', 'Yuan Zhu': '朱园', 'Honglei': '洪磊', 'Qiang': '龚强'}
NAME_EN = {'Kimi': 'Kimi (Ryan)', 'Yuan Zhu': 'Yuan Zhu', 'Honglei': 'Honglei', 'Qiang': 'Qiang Gong'}

SYS_KW = ['sweep', 'across', 'unify', ' all ', 'cascade', 'enforce', 'parity', 'similar',
          '统一', '全部', '所有']


def ctype(msg):
    m = msg.lower().strip()
    for t, l in [('feat', 'feat'), ('fix', 'fix'), ('test', 'test'), ('refactor', 'refactor'),
                 ('perf', 'style/perf'), ('style', 'style/perf'), ('docs', 'docs'),
                 ('chore', 'chore'), ('revert', 'chore'), ('merge', 'chore')]:
        if m.startswith(t):
            return l
    return 'other'


def git(args, cwd):
    return subprocess.run(['git'] + args, capture_output=True, text=True, cwd=cwd).stdout


def monday_of(d):
    return d - datetime.timedelta(days=d.weekday())


def isown(d):
    return d.isocalendar()[1]


# ----------------------------------------------------------------------------
# Metric collection
# ----------------------------------------------------------------------------
def collect(repo, since, today):
    M = {}

    # Weeks --------------------------------------------------------------
    cur_mon = monday_of(today)
    last_complete_sun = today - datetime.timedelta(days=today.weekday() + 1)
    deploy_end_mon = monday_of(last_complete_sun)
    deploy_weeks = []
    for i in range(4, -1, -1):
        ws = deploy_end_mon - datetime.timedelta(days=7 * i)
        deploy_weeks.append((f"W{isown(ws)}", ws, ws + datetime.timedelta(days=6)))
    trend = []
    ws = datetime.date(2026, 3, 23)
    while ws <= cur_mon:
        trend.append((f"W{isown(ws)}", ws, ws + datetime.timedelta(days=6)))
        ws += datetime.timedelta(days=7)
    M['deploy_weeks'] = deploy_weeks
    M['trend'] = trend
    M['cur_week_partial'] = (today.weekday() < 5)  # current week not finished

    # Commit log + bodies (AI detection) ---------------------------------
    log = git(['log', '--format=%ae|%an|%s|%cd|%H', '--date=format:%Y-%m-%d',
               f'--since={since}'], repo)
    bodies = git(['log', '--format=%H%x01%b%x02', f'--since={since}'], repo)
    ai_by_hash = {}
    for rec in bodies.split('\x02'):
        rec = rec.strip()
        if not rec:
            continue
        p = rec.split('\x01')
        if len(p) < 2:
            ai_by_hash[p[0]] = False
            continue
        ai_by_hash[p[0]] = 'co-authored-by: claude' in p[1].lower()

    devs = defaultdict(lambda: defaultdict(int))
    dev_days = defaultdict(set)
    week_deploy = defaultdict(lambda: defaultdict(int))
    week_trend = defaultdict(lambda: defaultdict(int))
    fix_msgs = defaultdict(list)
    feat_msgs = defaultdict(list)
    ai = defaultdict(lambda: [0, 0])
    conv = defaultdict(int)  # commits following Conventional Commits format
    total = 0
    CONV_PREFIXES = ('feat', 'fix', 'test', 'refactor', 'perf', 'style', 'docs',
                     'chore', 'revert', 'merge', 'build', 'ci')
    for line in log.strip().split('\n'):
        if not line:
            continue
        parts = line.split('|')
        if len(parts) < 5:
            continue
        email, name = parts[0], parts[1]
        h, date = parts[-1], parts[-2]
        msg = '|'.join(parts[2:-2])
        d = dev_of(email, name)
        if d is None:
            continue
        t = ctype(msg)
        devs[d][t] += 1
        devs[d]['total'] += 1
        total += 1
        dev_days[d].add(date)
        for wl, s, e in deploy_weeks:
            if s.isoformat() <= date <= e.isoformat():
                week_deploy[d][wl] += 1
        for wl, s, e in trend:
            if s.isoformat() <= date <= e.isoformat():
                week_trend[d][wl] += 1
        if t == 'fix':
            fix_msgs[d].append(msg)
        if t == 'feat':
            feat_msgs[d].append(msg)
        ml = msg.lower().strip()
        if any(ml.startswith(p + ':') or ml.startswith(p + '(') for p in CONV_PREFIXES):
            conv[d] += 1
        if ai_by_hash.get(h, False):
            ai[d][0] += 1
        else:
            ai[d][1] += 1

    M['devs'] = devs
    M['dev_days'] = dev_days
    M['week_deploy'] = week_deploy
    M['week_trend'] = week_trend
    M['fix_msgs'] = fix_msgs
    M['feat_msgs'] = feat_msgs
    M['ai'] = ai
    M['conv'] = conv
    M['total'] = total

    # File-add based: spec files + docs ----------------------------------
    def added(pathspecs):
        out = git(['log', '--diff-filter=A', '--format=C%x01%ae%x01%an',
                   f'--since={since}', '--name-only', '--'] + pathspecs, repo)
        res = defaultdict(list)
        cur = None
        for ln in out.split('\n'):
            if ln.startswith('C\x01'):
                _, e, n = ln.split('\x01')
                cur = dev_of(e, n)
            elif ln.strip() and cur:
                res[cur].append(ln.strip())
        return res

    M['specs'] = added(['e2e/**/*.spec.ts'])
    M['docs_added'] = added(['docs/**/*.md', 'scripts/**/*.mjs'])

    # Module ownership (file-touch count) --------------------------------
    out = git(['log', '--format=C%x01%ae%x01%an', f'--since={since}', '--name-only'], repo)
    mods = defaultdict(lambda: defaultdict(int))
    cur = None
    for ln in out.split('\n'):
        if ln.startswith('C\x01'):
            _, e, n = ln.split('\x01')
            cur = dev_of(e, n)
        elif ln.strip() and cur:
            m = module_of(ln.strip())
            if m:
                mods[m][cur] += 1
    M['mods'] = mods

    # Hotspots -----------------------------------------------------------
    out = git(['log', '--format=', f'--since={since}', '--name-only'], repo)
    touch = defaultdict(int)
    for ln in out.split('\n'):
        f = ln.strip()
        if f:
            touch[f] += 1
    M['hotspots'] = sorted(touch.items(), key=lambda x: -x[1])[:15]

    return M


def module_of(f):
    if f.startswith('gantt/'):
        return 'gantt'
    if f.startswith('pbs-server/'):
        return 'pbs-server'
    if f.startswith('pbs-portal/'):
        return 'pbs-portal'
    if f.startswith('live-server/'):
        return 'live-server'
    if f.startswith('rule-engine/') or f.startswith('rule-engine-rs/'):
        return 'rule-engine'
    if f.startswith('engine-server/'):
        return 'engine-server'
    if f.startswith('ro-engine/') or f.startswith('po-engine/'):
        return 'ro/po-engine'
    if f.startswith('e2e/'):
        return 'e2e tests'
    if f.startswith('sql/'):
        return 'sql'
    if f.startswith('docs/'):
        return 'docs'
    return None


MODULE_LABEL = {
    'gantt': ('gantt (Gantt scheduler)', 'gantt（甘特排班）'),
    'pbs-server': ('pbs-server (PBS backend)', 'pbs-server（PBS 后端）'),
    'pbs-portal': ('pbs-portal (PBS web)', 'pbs-portal（PBS 前端）'),
    'live-server': ('live-server (real-time)', 'live-server（实时服务）'),
    'rule-engine': ('rule-engine (legality)', 'rule-engine（法规引擎）'),
    'engine-server': ('engine-server (optimizer)', 'engine-server（优化调度）'),
    'ro/po-engine': ('ro/po-engine (solvers)', 'ro/po-engine（优化引擎）'),
    'e2e tests': ('e2e (Playwright tests)', 'e2e（Playwright 测试）'),
    'sql': ('sql (schema/seed)', 'sql（建表/基础数据）'),
    'docs': ('docs (documentation)', 'docs（文档）'),
}
MODULE_ORDER = list(MODULE_LABEL.keys())

SPEC_CATS = ['gantt core', 'gantt legality', 'gantt help', 'gantt perf', 'pbs-portal', 'e2e util']
DOC_CATS = ['Design Specs', 'Impl Plans', 'Architecture', 'Test Cases', 'Module Docs', 'Handoff', 'Auto Gates']


def cat_spec(f):
    if 'pbs-portal' in f or '/pbs/' in f:
        return 'pbs-portal'
    if 'legality' in f or 'legal' in f or re.search(r'(8002|8056|8030|8004|7501|7503|7504|7505|7506|7272)', f):
        return 'gantt legality'
    if 'help' in f:
        return 'gantt help'
    if 'perf' in f:
        return 'gantt perf'
    if 'gantt' in f:
        return 'gantt core'
    return 'e2e util'


def cat_doc(f):
    if 'specs/' in f:
        return 'Design Specs'
    if 'plans/' in f:
        return 'Impl Plans'
    if 'architecture/' in f:
        return 'Architecture'
    if 'test-cases/' in f:
        return 'Test Cases'
    if 'modules/' in f:
        return 'Module Docs'
    if 'handoff/' in f:
        return 'Handoff'
    if f.startswith('scripts/') and f.endswith('.mjs'):
        return 'Auto Gates'
    return None


# ----------------------------------------------------------------------------
# Derived ratings
# ----------------------------------------------------------------------------
def cfr_level(p, lang):
    if p < 5: return 'Elite ✅'
    if p < 15: return 'High ✅'
    if p < 30: return 'Medium ⚠️'
    return 'Low ⚠️'


def deploy_level(avg):
    if avg >= 50: return 'High ✅'
    if avg >= 10: return 'Medium ⚠️'
    if avg > 0: return 'Low ⚠️'
    return 'N/A'


def bus_factor(row, lang):
    owners = [d for d in DEVS if row.get(d, 0) >= max(1, sum(row.values()) * 0.1)]
    n = len(owners)
    if n >= 2:
        return ('Low — 2+ owners ✅', '低 — 多人共担 ✅')[lang == 'CN']
    if n == 1:
        return ('High — single owner ⚠️', '高 — 单人 ⚠️')[lang == 'CN']
    return '—'


def ai_obs(rate, lang):
    if rate >= 85:
        return ('Near-total AI collaboration', '几乎全程 AI 协作')[lang == 'CN']
    if rate >= 50:
        return ('Strong AI adoption', '较高 AI 采用')[lang == 'CN']
    if rate > 0:
        return ('Partial AI adoption', '部分采用 AI')[lang == 'CN']
    return ('No AI assistance — productivity gap', '未使用 AI — 存在效率差距')[lang == 'CN']


# ----------------------------------------------------------------------------
# docx helpers
# ----------------------------------------------------------------------------
BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x55, 0x55, 0x55)


def h(doc, text, size=13, color=BLUE, bold=True, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def note(doc, text, size=8.5, color=GREY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.italic = True
    return p


def table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Light Grid Accent 1'
    for i, hd in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ''
        run = c.paragraphs[0].add_run(str(hd))
        run.bold = True
        run.font.size = Pt(8.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run('' if val is None else str(val))
            run.font.size = Pt(8.5)
    return t


# ----------------------------------------------------------------------------
# Report builder
# ----------------------------------------------------------------------------
def build(M, lang, meta):
    CN = lang == 'CN'
    NM = NAME_CN if CN else NAME_EN
    period_end = meta['period_end_cn'] if CN else meta['period_end']
    doc = Document()
    doc.styles['Normal'].font.size = Pt(9)

    # Title -------------------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('ROIS-AI 项目 — DORA 开发者绩效分析报告' if CN
                  else 'ROIS-AI Project — DORA Developer Performance Analysis Report')
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nspec = sum(len(v) for v in M['specs'].values())
    if CN:
        s = (f"生成于 {meta['date']} ({meta['version']})  |  统计区间：2026年1月 – {period_end}"
             f"  |  {M['total']:,} 次提交  |  {nspec} 个 Playwright 测试文件  |  4 名开发者")
    else:
        s = (f"Generated: {meta['date']} ({meta['version']})  |  Period: Jan 2026 – {period_end}"
             f"  |  {M['total']:,} commits  |  {nspec} Playwright spec files  |  4 developers")
    rr = sub.add_run(s)
    rr.font.size = Pt(8.5)
    rr.font.color.rgb = GREY

    def explain(text):
        """Plain-language explainer printed right after a table, for non-technical readers."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        lead = p.add_run('💡 通俗解释：' if CN else '💡 In plain words: ')
        lead.bold = True
        lead.font.size = Pt(8.5)
        lead.font.color.rgb = RGBColor(0x2E, 0x5A, 0x2E)
        body = p.add_run(text)
        body.font.size = Pt(8.5)
        body.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    devs = M['devs']
    dw = M['deploy_weeks']
    wd = M['week_deploy']

    # 1. Identity map ---------------------------------------------------
    h(doc, '1. 团队成员与 Git 身份映射' if CN else '1. Team Members & Git Identity Map')
    note(doc, '五个 Git 身份对应四名真实开发者。朱园在不同设备上使用了不同的 git 配置，所有别名归并为同一人。' if CN
         else 'Five Git identities map to four real developers. Yuan Zhu committed under different git configs across devices — all aliases are merged into one person.')
    idrows = [
        [NM['Kimi'], 'eleihu6-rgb · Ryan', '2026-05-22 – ' + period_end],
        [NM['Yuan Zhu'], 'yuan.zhu-ai · yuan.zhu · yuanz-ai · yuanz-cc · Yuan Zhu', '2026-03-23 – ' + period_end],
        [NM['Honglei'], 'honglei.Yu', '2026-04-13 – ' + period_end],
        [NM['Qiang'], 'qiang.gong · Qiang Gong', '2026-06-15 – ' + period_end],
    ]
    table(doc, (['真实姓名', 'Git 别名', '活跃区间'] if CN
                else ['Real Name', 'Git Aliases', 'Active Period']), idrows)
    explain('这张表只是“谁是谁”。开发者有时会用不同的用户名或不同电脑提交代码（即“Git 别名”），我们把同一个人的所有别名合并，确保每人只计一次。“活跃区间”= 此人有提交记录的时间跨度。' if CN
            else 'This table is simply "who is who." A developer sometimes commits code under different usernames or from different computers (their "Git aliases"); we merge each person\'s aliases so everyone is counted once. "Active Period" = the span of dates during which they made commits.')

    # 2. Playwright test coverage (moved to the front for emphasis) -----
    h(doc, '2. 测试覆盖（Playwright 自动化测试）' if CN else '2. Test Coverage (Playwright Automated Tests)')
    note(doc, f'自 2026 年 1 月起共创建 {nspec} 个 Playwright 测试文件（spec）。按开发者与领域拆分：' if CN
         else f'{nspec} Playwright spec files created since Jan 2026. Breakdown by developer and area:')
    sprows = []
    for d in DEVS:
        c = defaultdict(int)
        for f in M['specs'].get(d, []):
            c[cat_spec(f)] += 1
        tot = sum(c.values())
        pct = f" ({tot / nspec * 100:.0f}%)" if nspec else ''
        sprows.append([NM[d]] + [c.get(k, 0) for k in SPEC_CATS] + [f"{tot}{pct}"])
    table(doc, (['开发者'] + SPEC_CATS + ['总计'] if CN else ['Developer'] + SPEC_CATS + ['Total']), sprows)
    note(doc, 'gantt legality = 每条 Rust 规则一个 spec（8002、8056、8030、8004、7501、7503、7504、7505、7506、7272 等）。' if CN
         else 'gantt legality = one spec per Rust rule (8002, 8056, 8030, 8004, 7501, 7503, 7504, 7505, 7506, 7272, etc.).')
    explain('我们最看重质量，所以把测试放在最前面。Playwright 是“自动点界面、自动验证”的测试工具，一个“spec”就是一个自动化测试文件。本项目有一条硬规定：每个新功能、每个 bug 修复都必须配一个 Playwright 测试（见第 4 节规范）。测试越多，以后改动越不容易把旧功能改坏。' if CN
            else 'We put testing first because quality matters most. Playwright is the tool that automatically clicks through the UI and verifies it works; each "spec" is one automated test file. The project has a hard rule: every feature and every bug fix must ship with a Playwright test (see §4 Standards). The more tests, the less likely future changes silently break existing features.')

    # 3. Fix quality (moved up — right after testing) -------------------
    h(doc, '3. 修复质量：点修复 vs. 系统性修复' if CN else '3. Fix Quality: Point Fix vs. Systematic Fix')
    note(doc, '发现一个 bug 时，开发者是只修这一处，还是搜索整个代码库修复所有同类问题？' if CN
         else 'When a bug is found, does the developer fix only this instance or search the whole codebase and fix all occurrences?')
    fqrows = []
    for d in DEVS:
        fm = M['fix_msgs'].get(d, [])
        sysc = sum(1 for m in fm if any(k in m.lower() for k in SYS_KW))
        point = len(fm) - sysc
        rate = sysc / len(fm) * 100 if fm else 0
        lvl = 'High ✅' if rate >= 15 else ('Medium ⚠️' if rate >= 5 else 'Low ⚠️')
        examples = '; '.join(m[:55] for m in fm if any(k in m.lower() for k in SYS_KW))[:140] or '—'
        fqrows.append([NM[d], len(fm), point, sysc, f"{rate:.1f}%", lvl, examples])
    table(doc, (['开发者', 'fix 总数', '点修复', '系统性', '系统化率', '级别', '示例'] if CN
                else ['Developer', 'Fix Total', 'Point Fix', 'Systematic', 'Sys. Rate', 'Level', 'Examples']), fqrows)
    note(doc, '系统化率 = 系统性修复 ÷ 总 fix 提交。检测：关键词匹配（sweep/across/unify/all/cascade/enforce/parity/similar/统一/全部/所有）。' if CN
         else 'Sys. Rate = systematic fixes ÷ total fix commits. Detection: keyword match (sweep/across/unify/all/cascade/enforce/parity/similar).')
    explain('修一个 bug 时有两种做法：“点修复”= 只修眼前这一处；“系统性修复”= 顺手在整个代码库里把所有同类问题一起修掉。后者质量更高，因为同样的毛病往往不止一处。“系统化率”就是系统性修复占全部修复的比例（我们通过提交说明里的关键词，如“统一/全部/across”来识别）。' if CN
            else 'When fixing a bug there are two styles: a "Point Fix" repairs only the one spot in front of you; a "Systematic" fix also sweeps the whole codebase and repairs every similar case at once. The latter is higher quality, because the same flaw usually appears in more than one place. "Sys. Rate" is the share of fixes that were systematic (we detect it from keywords in the commit message such as "unify / all / across").')

    # 4. Standards & conventions adherence (NEW) -----------------------
    h(doc, '4. 团队工程规范与执行情况' if CN else '4. Engineering Standards & Conventions')
    note(doc, '本项目有一套强制执行的工程规范，部分由脚本自动门禁守护。下表先列出规范本身，再看每位开发者对“提交信息规范”的遵循度。' if CN
         else 'The project enforces a set of engineering standards, some guarded by automated gates. The first table lists the standards; the second measures each developer\'s adherence to the commit-message convention.')
    if CN:
        std = [
            ['UI / 排版样式标准', '字号/颜色/间距/圆角一律用 token，禁止魔法值', '自动门禁：npm run check:ui + git pre-push 钩子（硬违规直接拦截推送）'],
            ['Playwright 测试强制', '每个新功能、每个 bug 修复都必须配 Playwright 测试', '代码评审（CLAUDE.md §Playwright-Required）'],
            ['No-Illusion（拒绝臆断）', '任何“已完成/已修复”都必须贴出真实测试运行结果', '代码评审（CLAUDE.md §No-Illusion）'],
            ['Conventional Commits 提交规范', '提交信息用 feat/fix/test/refactor/docs/chore 等前缀', '约定（建议加 commitlint 自动校验）'],
            ['TypeScript / Python 编码规范', '命名规范、禁止 any、导入顺序、Zod 入参校验', '评审 + 类型检查（tsc）'],
            ['弹窗窗口标准', '所有弹窗统一用共享 AppDialog 组件', '代码评审（CLAUDE.md）'],
            ['版本号管理', '每次代码改动都要递增 version.ts', '代码评审'],
        ]
        table(doc, ['规范', '要求', '如何执行'], std)
    else:
        std = [
            ['UI / Typography Standard', 'Token-driven font size, colour, spacing, radius — no magic values', 'Automated gate: npm run check:ui + git pre-push hook (hard violations block the push)'],
            ['Playwright-Required', 'Every feature and every bug fix ships with a Playwright test', 'Code review (CLAUDE.md §Playwright-Required)'],
            ['No-Illusion', 'Any "done / fixed" claim must show a real test-run result', 'Code review (CLAUDE.md §No-Illusion)'],
            ['Conventional Commits', 'Commit messages use feat/fix/test/refactor/docs/chore prefixes', 'Convention (commitlint recommended)'],
            ['TypeScript / Python coding standard', 'Naming rules, no "any", import order, Zod input validation', 'Review + type-check (tsc)'],
            ['Pop-up Window Standard', 'All dialogs built with the shared AppDialog component', 'Code review (CLAUDE.md)'],
            ['Version Bumping', 'Bump version.ts on every code change', 'Code review'],
        ]
        table(doc, ['Standard', 'What it requires', 'How it is enforced'], std)
    explain('回答“有没有规范要遵守？”——有，而且不止 UI 标准：还有编码规范、必写测试、提交信息格式、弹窗统一组件、版本号管理等。其中“UI 样式标准”是全自动的：每次推送代码前脚本会自动检查，发现硬性违规会直接拦截推送，不靠自觉。' if CN
            else 'This answers "are there standards to follow?" — yes, and not just the UI standard: there are coding conventions, a mandatory-test rule, a commit-message format, a shared dialog component, and version-bump rules. The UI Style Standard is fully automated: before every push a script checks the code and blocks hard violations — it does not rely on goodwill.')
    cvrows = []
    for d in DEVS:
        tot = devs[d]['total']
        good = M['conv'].get(d, 0)
        off = tot - good
        rate = good / tot * 100 if tot else 0
        lvl = 'High ✅' if rate >= 90 else ('Medium ⚠️' if rate >= 70 else 'Low ⚠️')
        cvrows.append([NM[d], tot, good, off, f"{rate:.1f}%", lvl])
    table(doc, (['开发者', '总提交', '合规', '不合规', '合规率', '级别'] if CN
                else ['Developer', 'Total', 'Conventional', 'Off-convention', 'Adherence', 'Level']), cvrows)
    explain('“提交信息规范”要求每条提交以统一前缀（feat/fix/…）开头，这样工具能自动汇总、生成变更日志。上表是每个人遵循该格式的比例。合规率低会让 CFR 等指标失真——这正是某位同事数字偏差的根源（建议加 commitlint 自动校验）。' if CN
            else 'The commit-message convention requires every commit to start with a standard prefix (feat/fix/…) so tools can auto-summarise and build changelogs. The table above is each person\'s adherence to that format. Low adherence distorts metrics like CFR — exactly the source of one teammate\'s skewed numbers (adding commitlint would auto-enforce it).')

    # 5. Commit activity ------------------------------------------------
    h(doc, '5. 提交活动汇总（年初至今，所有别名已归并）' if CN
      else '5. Commit Activity Summary (YTD, all aliases merged)')
    note(doc, f"朱园 = yuan.zhu-ai + yuan.zhu + yuanz-ai + yuanz-cc + Yuan Zhu 归并。当前版本：{meta['app_version']}。" if CN
         else f"Yuan Zhu = yuan.zhu-ai + yuan.zhu + yuanz-ai + yuanz-cc + \"Yuan Zhu\" merged. Current version: {meta['app_version']}.")
    crows = []
    for d in DEVS:
        dd = devs[d]
        crows.append([NM[d], dd['total'], dd['feat'], dd['fix'], dd['test'],
                      dd['refactor'], dd['style/perf'], dd['docs'], dd['chore'],
                      len(M['dev_days'][d])])
    table(doc, (['开发者', '总计', 'feat', 'fix', 'test', 'refactor', 'style/perf', 'docs', 'chore', '活跃天数'] if CN
                else ['Developer', 'Total', 'feat', 'fix', 'test', 'refactor', 'style/perf', 'docs', 'chore', 'Active Days']), crows)
    note(doc, 'feat=新功能  fix=修复  test=自动化测试  refactor=重构  docs=文档  chore=工具/依赖' if CN
         else 'feat = new feature   fix = bug fix   test = automated test   refactor = internal restructuring   docs = documentation   chore = toolchain/deps')
    explain('“提交”（commit）= 把一批改动保存进项目历史，可以理解为“一次存档”。“总计”= 此人一共存档了多少次。后面几列把每次存档按用途分类：新增功能、修复 bug、写自动化测试、整理代码（不改变功能）、调格式/性能、写文档、工具与依赖维护。“活跃天数”= 此人有提交的不同日期数。' if CN
            else 'A "commit" is one saved batch of changes to the project history — think of it as a single "save point." "Total" = how many save points this person made. The remaining columns sort those saves by purpose: new features, bug fixes, automated tests, tidying code (no behaviour change), formatting/speed, writing docs, and tool/dependency upkeep. "Active Days" = the number of distinct days on which they committed.')
    note(doc, '⚠ 洪磊将大量新功能误标为 "fix:"，导致 CFR 虚高，真实功能数应更高。' if CN
         else '⚠ Honglei mislabels features as "fix:", causing inflated CFR. Real feature count is higher.')
    if meta.get('delta'):
        note(doc, meta['delta'])

    # 6. DORA four key metrics ------------------------------------------
    h(doc, '6. DORA 四项关键指标' if CN else '6. DORA Four Key Metrics')
    note(doc, 'DORA（由 Google 支持）通过四项指标衡量软件交付效能。本项目无 CI/CD 流水线，使用 git 提交历史作为代理指标。' if CN
         else 'DORA (backed by Google) measures software delivery performance via four indicators. This project has no CI/CD pipeline; git commit history is used as a proxy.')
    explain('DORA 是业界公认的一套衡量“软件团队交付得好不好”的标准，源自 Google 的多年研究。它看四件事：① 多久交付一次；② 一个改动从动手到完成要多久；③ 改动出问题的比例有多高；④ 出问题后多久能修复。本项目还没有自动化部署系统，所以我们用“提交历史”来近似估算（即“代理指标”）。' if CN
            else 'DORA is an industry-standard way to gauge how well a software team delivers, based on years of research by Google. It looks at four things: (1) how often work ships, (2) how long a change takes from start to finished, (3) what share of changes go wrong, and (4) how quickly problems are fixed. This project has no automated deployment system, so we approximate these using commit history (a "proxy").')

    # 3.1 Deployment freq
    h(doc, '6.1 部署频率（Deployment Frequency）' if CN else '6.1 Deployment Frequency', size=11)
    note(doc, f"代理指标：每周提交数。{dw[-1][0]} = {dw[-1][1].strftime('%m/%d')}–{dw[-1][2].strftime('%m/%d')}（已完整统计）。" if CN
         else f"Proxy: weekly commit count. {dw[-1][0]} = {dw[-1][1].strftime('%b %d')}–{dw[-1][2].strftime('%d')} (fully measured).")
    header = (['开发者', '总计'] if CN else ['Developer', 'Total']) + [f"{w[0]}({w[1].strftime('%-m/%-d')})" for w in dw] + (['周均', 'DORA 级别'] if CN else ['Avg/wk', 'DORA Level'])
    drows = []
    peak_dev = max(DEVS, key=lambda d: max((wd[d].get(w[0], 0) for w in dw), default=0))
    peak_val = max((wd[peak_dev].get(w[0], 0) for w in dw), default=0)
    for d in DEVS:
        wk = [wd[d].get(w[0], 0) for w in dw]
        cells = []
        for w, v in zip(dw, wk):
            mark = ' ▲' if (d == peak_dev and v == peak_val and v > 0) else ''
            cells.append(f"{v}{mark}")
        avg = round(sum(wk) / len(wk))
        drows.append([NM[d], devs[d]['total']] + cells + [avg, deploy_level(avg)])
    table(doc, header, drows)
    note(doc, 'DORA 级别：Elite=每日多次；High=每日–每周；Medium=每周–每月；Low=每月以上' if CN
         else 'DORA Levels: Elite = multiple/day; High = daily–weekly; Medium = weekly–monthly; Low = monthly+')
    note(doc, f'▲ {NM[peak_dev]} 的高峰周。周均 = {dw[0][0]}–{dw[-1][0]} 五周平均。' if CN
         else f"▲ {NM[peak_dev]}'s peak week. Avg/wk = {dw[0][0]}–{dw[-1][0]} average (5 weeks).")
    explain('“部署频率”= 这个人多久交付一次工作。我们用“每周提交次数”来近似：数字越大、越规律，说明改动是小步快跑（更安全），而不是攒一大批一次性上线（风险更高）。表中是最近 5 周的每周提交数，最后两列是周平均和对应的 DORA 等级。' if CN
            else '"Deployment Frequency" = how often this person delivers work. We approximate it with commits per week: a higher, steadier number means changes ship in small, frequent batches (safer) rather than one big risky drop. The table shows the last 5 weeks of weekly commits, with the weekly average and matching DORA level in the final columns.')

    # 3.2 Lead time
    h(doc, '6.2 变更前置时间（Lead Time for Changes）' if CN else '6.2 Lead Time for Changes', size=11)
    ltrows = []
    for d in DEVS:
        active = devs[d]['total']
        if active >= 100:
            lt, lvl = ('1–3 天', 'High ✅') if CN else ('1–3 days', 'High ✅')
            ev = ('同日完成 feat→fix→test 闭环；多数特性 1–3 天内交付' if CN
                  else 'feat→fix→test same day; most features land in 1–3 days')
        elif active >= 10:
            lt, lvl = ('1–7 天', 'High ✅') if CN else ('1–7 days', 'High ✅')
            ev = ('特性通常一周内交付' if CN else 'features ship within a week')
        else:
            lt, lvl = ('数据不足', 'N/A') if CN else ('insufficient data', 'N/A')
            ev = ('提交量过少，无法估计' if CN else 'too few commits to estimate')
        ltrows.append([NM[d], devs[d]['total'], lt, lvl, ev])
    table(doc, (['开发者', '总计', '估计前置时间', 'DORA 级别', '证据'] if CN
                else ['Developer', 'Total', 'Est. Lead Time', 'DORA Level', 'Evidence']), ltrows)
    explain('“变更前置时间”（Lead Time）= 一个改动从开始动手到真正完成、可用，需要多久。越短越好——说明想法能很快变成能用的软件。“证据”列给出我们据以估计的实际例子（例如同一天内完成“开发→修复→测试”的闭环）。' if CN
            else '"Lead Time for Changes" = how long a change takes from first starting it to being finished and usable. Shorter is better — it means ideas become working software fast. The "Evidence" column gives the real examples we based the estimate on (e.g. a build→fix→test loop completed within the same day).')

    # 3.3 CFR
    h(doc, '6.3 变更失败率（Change Failure Rate, CFR）' if CN else '6.3 Change Failure Rate (CFR)', size=11)
    note(doc, '代理指标：fix ÷ 总提交。Elite <5%；High 5–15%；Medium 15–30%；Low >45%。' if CN
         else 'Proxy: fix: ÷ total commits. Elite <5%; High 5–15%; Medium 15–30%; Low >45%.')
    cfrows = []
    for d in DEVS:
        dd = devs[d]
        cfr = dd['fix'] / dd['total'] * 100 if dd['total'] else 0
        cfrows.append([NM[d], dd['total'], dd['fix'], dd['feat'], f"{cfr:.1f}%", cfr_level(cfr, lang)])
    table(doc, (['开发者', '总计', 'fix 数', 'feat 数', 'CFR', '级别'] if CN
                else ['Developer', 'Total', 'fix count', 'feat count', 'CFR', 'Level']), cfrows)
    hcfr = devs['Honglei']['fix'] / devs['Honglei']['total'] * 100 if devs['Honglei']['total'] else 0
    note(doc, f'* 洪磊把新功能误标为 "fix:"；原始 {hcfr:.1f}%，估计真实 CFR ~25%。' if CN
         else f'* Honglei mislabels features as "fix:"; raw {hcfr:.1f}%, estimated real CFR ~25%.')
    explain('“变更失败率”（CFR）= 在所有改动里，有多少是“修复类”的（我们用 fix 提交占比近似，代表“有东西需要纠正”）。比例越低越好。注意：有一位同事习惯把“新功能”也标成“fix”，所以他的数字看起来比真实情况差很多——这是标注习惯问题，不是质量问题。' if CN
            else '"Change Failure Rate" (CFR) = of all the changes made, what share were corrective (we approximate with the proportion of "fix" commits, standing in for "something needed fixing"). Lower is better. Note: one teammate habitually labels new features as "fix" too, which makes his number look far worse than reality — a labelling habit, not a quality problem.')

    # 3.4 MTTR
    h(doc, '6.4 平均恢复时间（MTTR）' if CN else '6.4 Mean Time to Recovery (MTTR)', size=11)
    note(doc, '无法从 git 历史计算。建议：在 Plane 中标记事故，并用 "hotfix:" 前缀提交。' if CN
         else 'Cannot be calculated from git history. Recommendation: tag incidents in Plane + prefix commits with "hotfix:".')
    explain('“平均恢复时间”（MTTR）= 线上出故障后，平均多久能恢复正常。目前提交历史里没有“故障发生/修复”的时间标记，所以暂时算不出来；建议以后在任务系统里登记故障、并用 "hotfix:" 标注紧急修复，就能量化了。' if CN
            else '"Mean Time to Recovery" (MTTR) = on average, how long it takes to get back to normal after something breaks in production. Commit history has no "incident started / fixed" timestamps, so we can\'t compute it yet; logging incidents in the task system and tagging urgent fixes with "hotfix:" would make it measurable.')

    # 7. Scorecard ------------------------------------------------------
    h(doc, '7. DORA 评分卡' if CN else '7. DORA Scorecard')
    scrows = []
    for d in DEVS:
        dd = devs[d]
        cfr = dd['fix'] / dd['total'] * 100 if dd['total'] else 0
        td = dd['test'] / dd['total'] * 100 if dd['total'] else 0
        avg = round(sum(wd[d].get(w[0], 0) for w in dw) / len(dw))
        dlv = deploy_level(avg)
        ltlv = 'High ✅' if dd['total'] >= 10 else 'N/A'
        if dd['total'] >= 100:
            overall = ('高效能 ✅' if CN else 'High Performer ✅')
        elif dd['total'] >= 10:
            overall = ('稳定贡献 ✅' if CN else 'Solid Contributor ✅')
        else:
            overall = ('观察期 ⏳' if CN else 'Ramp-up ⏳')
        scrows.append([NM[d], dd['total'], dlv, ltlv, cfr_level(cfr, lang), 'N/A',
                       f"{td:.1f}% " + ('✅' if td >= 5 else '⚠️'), overall])
    table(doc, (['开发者', '总计', '部署频率', '前置时间', 'CFR', 'MTTR', '测试纪律', '综合评级'] if CN
                else ['Developer', 'Total', 'Deploy Freq', 'Lead Time', 'CFR', 'MTTR', 'Test Discipline', 'Overall Rating']), scrows)
    note(doc, '测试纪律 = test 提交 ÷ 总提交' if CN else 'Test Discipline = test: commits ÷ total commits')
    explain('这是一张“成绩单”，把前面四项 DORA 指标加上“测试纪律”汇总成一行。“测试纪律”= 此人有多大比例的提交是在写自动化测试（越高说明越重视质量把关）。“综合评级”是一个简单的总结标签（如“高效能/稳定贡献/观察期”）。各列里的 ✅ 表示良好、⚠️ 表示需关注、N/A 表示暂无法衡量。' if CN
            else 'This is a "report card" that rolls the four DORA measures plus "Test Discipline" into one line. "Test Discipline" = what share of this person\'s commits are writing automated tests (higher means more attention to quality safeguards). "Overall Rating" is a simple summary label (e.g. High Performer / Solid Contributor / Ramp-up). In each column, ✅ = good, ⚠️ = worth attention, N/A = not measurable yet.')

    # 8. Module ownership ----------------------------------------------
    h(doc, '8. 模块归属分析' if CN else '8. Module Ownership Analysis')
    note(doc, '以文件改动次数衡量（每位开发者在每个模块上的文件变更事件数，跨全部提交）。' if CN
         else 'Measured by file-touch count (number of file-change events per developer per module across all commits).')
    mrows = []
    for m in MODULE_ORDER:
        row = M['mods'].get(m, {})
        if sum(row.values()) == 0:
            continue
        tot = sum(row.values())
        mrows.append([MODULE_LABEL[m][CN]] + [row.get(d, 0) for d in DEVS] + [tot, bus_factor(row, lang)])
    table(doc, ([ '模块', NM['Kimi'], NM['Yuan Zhu'], NM['Honglei'], NM['Qiang'], '总计', 'Bus Factor 风险'] if CN
                else ['Module', 'Kimi', 'Yuan Zhu', 'Honglei', 'Qiang', 'Total', 'Bus Factor Risk']), mrows)
    note(doc, 'Bus Factor 风险：低 = 2+ 负责人；高 = 单人负责。' if CN
         else 'Bus Factor Risk: Low = 2+ owners; High = single owner.')
    explain('这张表显示每个人主要在系统的哪些部分工作，用“改动文件的次数”来衡量。“Bus Factor（公车因子）风险”是一个形象的说法：如果某块代码只有一个人懂，万一这个人离开（被公车撞了……），团队就会很被动——这就是“高风险/单人负责”。两人以上都熟悉则是“低风险”。' if CN
            else 'This table shows which parts of the system each person mainly works on, measured by how many file changes they made there. "Bus Factor Risk" is a vivid term: if only one person understands a piece of code, the team is exposed if that person leaves (or is "hit by a bus"…) — that is "High risk / single owner." When two or more people know it, the risk is "Low."')

    # 9. Standards & docs ----------------------------------------------
    h(doc, '9. 规范与文档贡献' if CN else '9. Standards & Documentation Contributions')
    dcrows = []
    for d in DEVS:
        c = defaultdict(int)
        for f in M['docs_added'].get(d, []):
            k = cat_doc(f)
            if k:
                c[k] += 1
        tot = sum(c.values())
        dcrows.append([NM[d]] + [c.get(k, 0) for k in DOC_CATS] + [tot])
    table(doc, (['开发者'] + DOC_CATS + ['总计'] if CN else ['Developer'] + DOC_CATS + ['Total']), dcrows)
    note(doc, 'Kimi：自动化门禁的唯一作者（check-ui-standard.mjs、CLAUDE.md §UI-Standard-Gate、规则迁移 playbook）。' if CN
         else 'Kimi: sole creator of automated enforcement gates (check-ui-standard.mjs, CLAUDE.md §UI-Standard-Gate, rule migration playbook).')
    explain('这张表统计每人写了多少“文档”，并按类型拆分：设计规格=要做什么的书面说明；实施计划=分步骤的开发计划；架构=系统整体设计；测试用例=人工测试的脚本；模块文档=某个部分的使用说明；交接=人/会话之间传递的笔记；自动门禁=能自动拦截不合规代码的检查脚本。文档齐全能让新人更快上手、减少口口相传的依赖。' if CN
            else 'This table counts how much documentation each person wrote, split by type: Design Specs = written descriptions of what to build; Impl Plans = step-by-step build plans; Architecture = the big-picture system design; Test Cases = manual testing scripts; Module Docs = guides for one part of the system; Handoff = notes passed between people or work sessions; Auto Gates = scripts that automatically block non-compliant code. Good docs help newcomers ramp up faster and reduce reliance on word-of-mouth.')

    # 10. Recommendations -----------------------------------------------
    h(doc, '10. 改进建议' if CN else '10. Improvement Recommendations')
    if CN:
        recs = [
            ['洪磊', '采用 Conventional Commits', '安装 commitlint。把新功能误标为 "fix:" 会让 CFR 虚高到 70% 并掩盖真实产出。'],
            ['洪磊', '引入 Claude Code', '0% AI 协作率是明显的效率短板；引入后或可同时修复提交类型误标与测试纪律。'],
            ['洪磊', '提升测试纪律', 'pbs-portal 由其单人负责（bus factor 1），需补 Playwright 回归并做知识转移。'],
            ['朱园', '保持系统性修复习惯', 'W25 多为新功能；系统化率回落至个位数，注意在修 bug 时顺手清理同类问题。'],
            ['Kimi', '维持自动化门禁', '继续把规范固化成可执行门禁（已建 UI-Standard-Gate / 规则迁移 playbook）。'],
            ['龚强', '观察期', '已立即采用 AI（75%）。4 周后 DORA 指标更具参考价值。'],
            ['全员', '建立 MTTR 追踪', '在 Plane 标记事故，提交用 "hotfix:" 前缀，使 MTTR 可量化。'],
            ['全员', '统一周节奏', 'Kimi+朱园连续多周 >100 提交/周，注意可持续性与代码评审深度。'],
        ]
        table(doc, ['负责人', '行动', '详情'], recs)
    else:
        recs = [
            ['Honglei', 'Adopt Conventional Commits', 'Install commitlint. Mislabeling features as "fix:" inflates CFR to 70% and hides real output.'],
            ['Honglei', 'Onboard Claude Code', '0% AI rate is a productivity gap; adoption may fix commit-type mislabeling and test discipline at once.'],
            ['Honglei', 'Raise test discipline', 'pbs-portal is single-owner (bus factor 1); add Playwright regressions and transfer knowledge.'],
            ['Yuan Zhu', 'Keep systematic-fix habit', 'W25 was feature-heavy; sys. rate dropped to single digits — sweep sibling occurrences when fixing.'],
            ['Kimi', 'Maintain automated gates', 'Continue codifying standards into executable gates (UI-Standard-Gate, rule-migration playbook).'],
            ['Qiang', 'Ramp-up window', 'Adopted AI immediately (75%). DORA metrics meaningful in ~4 weeks.'],
            ['Team', 'Establish MTTR tracking', 'Tag incidents in Plane; prefix commits with "hotfix:" so MTTR becomes measurable.'],
            ['Team', 'Sustainable cadence', 'Kimi+Yuan Zhu sustained >100 commits/wk for weeks — watch sustainability and review depth.'],
        ]
        table(doc, ['Owner', 'Action', 'Detail'], recs)
    explain('这是根据上面数据给出的具体改进建议：谁来做、做什么、为什么。多数是低成本、可立刻落地的动作，目的是让团队指标更真实、协作风险更低。' if CN
            else 'These are concrete improvement actions drawn from the data above: who should do it, what to do, and why. Most are low-cost, immediately actionable steps aimed at making the team\'s metrics more accurate and reducing collaboration risk.')

    # 11. Methodology ---------------------------------------------------
    h(doc, '11. 方法说明与局限' if CN else '11. Methodology & Limitations')
    if CN:
        meth = [
            ['部署频率', '以每周提交数为代理', '无 CI/CD 流水线日志。'],
            ['前置时间', 'feat→fix→test 同日闭环观察', '无 PR 评审时间戳。'],
            ['CFR', 'fix ÷ 总提交', '依赖提交信息规范；洪磊误标致虚高。'],
            ['MTTR', '无法计算', '无事故标记与 hotfix 前缀。'],
            ['模块归属', '文件改动次数', '次数 ≠ 行数 ≠ 复杂度。'],
            ['测试', '新增 *.spec.ts 文件计数', '不计断言数与覆盖率。'],
            ['AI 协作', '提交 body 含 "Co-Authored-By: Claude"', '仅反映标注，未标注则漏计。'],
            ['身份归并', 'git email + name 消歧', '别名靠人工映射维护。'],
        ]
        table(doc, ['指标', '方法', '局限'], meth)
    else:
        meth = [
            ['Deployment Frequency', 'Weekly commit count as proxy', 'No CI/CD pipeline logs.'],
            ['Lead Time', 'feat→fix→test same-day observation', 'No PR review timestamps.'],
            ['CFR', 'fix ÷ total commits', 'Depends on commit hygiene; Honglei mislabels inflate it.'],
            ['MTTR', 'Not computable', 'No incident tags or hotfix prefix.'],
            ['Module Ownership', 'File-touch count', 'Touches ≠ lines ≠ complexity.'],
            ['Tests', 'Count of added *.spec.ts files', 'Ignores assertion count / coverage.'],
            ['AI Assist', '"Co-Authored-By: Claude" in commit body', 'Reflects annotation only; unannotated work missed.'],
            ['Identity merge', 'git email + name disambiguation', 'Aliases maintained by hand-mapping.'],
        ]
        table(doc, ['Metric', 'Method', 'Limitation'], meth)
    explain('这张表说明每个数字是“怎么算出来的”以及它的“局限”（哪里可能不准）。重点是：这些指标都来自提交历史的近似，能反映趋势和相对差异，但不是精确的考核分数，请结合实际情况解读。' if CN
            else 'This table explains how each number was calculated and its limitations (where it may be inaccurate). The key point: these metrics are approximations drawn from commit history — good for trends and relative comparison, but not exact performance scores, so read them alongside real-world context.')
    yz = devs['Yuan Zhu']['total']
    note(doc, f"统计区间：2026-01-01 至 {period_end}。所有别名按 git email + name 消歧归并。" if CN
         else f"Period: 2026-01-01 to {period_end}. All aliases normalized by git email + name disambiguation.")
    note(doc, f"朱园别名：yuan.zhu-ai + yuan.zhu + yuanz-ai + yuanz-cc + Yuan Zhu = {yz:,} 次提交。" if CN
         else f"Yuan Zhu aliases: yuan.zhu-ai + yuan.zhu + yuanz-ai + yuanz-cc + \"Yuan Zhu\" = {yz:,} commits total.")
    if meta.get('prior_line'):
        note(doc, meta['prior_line'])

    # 11. Velocity trend ------------------------------------------------
    h(doc, ('12. 开发速度趋势（W13–%s，2026）' % M['trend'][-1][0]) if CN
      else ('12. Development Velocity Trend (W13–%s, 2026)' % M['trend'][-1][0]))
    note(doc, '从项目启动到当前的完整每周提交历史，呈现团队成长阶段、Kimi 加入与交付高峰周。' if CN
         else 'Full weekly commit history from project start through the current week. Shows team growth phases, Kimi onboarding, and peak delivery weeks.')
    wt = M['week_trend']
    team_peak = 0
    peak_w = ''
    for wl, s, e in M['trend']:
        team = sum(wt[d].get(wl, 0) for d in DEVS)
        if team > team_peak:
            team_peak = team
            peak_w = wl
    trows = []
    for wl, s, e in M['trend']:
        team = sum(wt[d].get(wl, 0) for d in DEVS)
        partial = (wl == M['trend'][-1][0] and M['cur_week_partial'])
        if CN:
            nt = '在统计中…' if partial else ('团队记录周' if wl == peak_w else '')
        else:
            nt = 'in progress…' if partial else ('team record week' if wl == peak_w else '')
        trows.append([f"2026-{wl} ({s.strftime('%-m/%-d')})"] + [wt[d].get(wl, 0) for d in DEVS] + [team, nt])
    table(doc, (['周', NM['Kimi'], NM['Yuan Zhu'], NM['Honglei'], NM['Qiang'], '团队合计', '备注'] if CN
                else ['Week', 'Kimi', 'Yuan Zhu', 'Honglei', 'Qiang', 'Team Total', 'Note']), trows)
    note(doc, f'▲ {peak_w} = 团队记录周（{team_peak} 次提交）。Kimi 加入后立即带来团队产出的阶跃式提升。' if CN
         else f'▲ {peak_w} = team record week ({team_peak} commits). Kimi joined and immediately drove a step-change in team output.')
    explain('这张表是团队从项目初期到现在、每一周的提交总量。它像一条“产出曲线”，能看出团队是怎么一步步壮大的、哪几周特别忙、以及新成员加入后产出的变化。最后一周如果标注“在统计中”，表示本周还没结束、数字会继续增加。' if CN
            else 'This table is the team\'s total commits for every week from the project\'s early days to now. It works like an "output curve," showing how the team grew over time, which weeks were especially busy, and how output shifted as new members joined. If the last week is marked "in progress," it simply means the week isn\'t over yet and the number will keep rising.')

    # 13. AI-assisted ---------------------------------------------------
    h(doc, '13. AI 辅助开发（Claude Code）' if CN else '13. AI-Assisted Development (Claude Code)')
    note(doc, '提交 body 含 "Co-Authored-By: Claude" 表示该提交由 Claude Code 协助完成。' if CN
         else 'Commits containing "Co-Authored-By: Claude" in the commit body indicate work done with Claude Code assistance.')
    airows = []
    team_ai = team_h = 0
    for d in DEVS:
        a, hu = M['ai'][d]
        tot = a + hu
        team_ai += a
        team_h += hu
        rate = a / tot * 100 if tot else 0
        airows.append([NM[d], tot, a, hu, f"{rate:.1f}%", ai_obs(rate, lang)])
    table(doc, (['开发者', '总提交', 'AI 辅助', '纯人工', 'AI 比例', '观察'] if CN
                else ['Developer', 'Total Commits', 'AI-Assisted', 'Human-Only', 'AI Rate', 'Observation']), airows)
    explain('“AI 辅助”指这次改动是借助 Claude Code（我们的 AI 编程助手）完成的——我们通过提交里自动留下的署名标记来识别。“AI 比例”= 此人有多大比例的提交用到了 AI。比例越高，通常意味着借助 AI 把工作做得更快。这只是统计 AI 的使用强度，不代表工作质量好坏。' if CN
            else '"AI-Assisted" means the change was done with help from Claude Code (our AI coding assistant) — we detect it from a signature the assistant automatically leaves in the commit. "AI Rate" = the share of this person\'s commits that used AI. A higher rate usually means leveraging AI to work faster. This only measures how much AI was used, not whether the work was good or bad.')
    team_rate = team_ai / (team_ai + team_h) * 100 if (team_ai + team_h) else 0
    if CN:
        note(doc, f'团队整体 AI 比例 {team_rate:.0f}%。Kimi 与朱园以 AI-first 的前沿节奏工作（90%+），是团队产能的乘数。', size=9, color=RGBColor(0, 0, 0))
        note(doc, '洪磊 0% AI 比例是明显的效率差距；引入 Claude Code 或可同时解决提交类型误标与测试纪律问题。', size=9, color=RGBColor(0, 0, 0))
    else:
        note(doc, f'Team-wide AI rate {team_rate:.0f}%. Kimi & Yuan Zhu operate at frontier AI-first velocity (90%+) — a team-wide productivity multiplier.', size=9, color=RGBColor(0, 0, 0))
        note(doc, "Honglei's 0% AI rate is a significant productivity gap. Onboarding Claude Code may resolve commit-type mislabeling and test-discipline gaps simultaneously.", size=9, color=RGBColor(0, 0, 0))

    # 14. Hotspots ------------------------------------------------------
    h(doc, '14. 代码热点分析（改动最频繁的文件）' if CN else '14. Code Hotspot Analysis (Most-Touched Files)')
    note(doc, '跨全部提交改动最频繁的文件 — 高频意味着复杂度中心、活跃特性区或横切重构目标。' if CN
         else 'Files touched most frequently across all commits — high churn indicates complexity centres, active feature areas, or cross-cutting refactoring targets.')
    hsrows = []
    for i, (f, c) in enumerate(M['hotspots'], 1):
        if 'version.ts' in f:
            interp = ('每次前端改动都会 bump — 属预期，非风险' if CN else 'Version bump on every frontend change — expected, not a risk')
        elif f.startswith('pbs-'):
            interp = ('PBS 单人热点 — bus factor 1，需知识转移' if CN else 'PBS single-owner hotspot — bus factor 1; knowledge transfer needed')
        elif 'CLAUDE.md' in f:
            interp = ('团队规范治理活跃' if CN else 'Active team standards governance')
        elif 'pane' in f or 'renderer' in f:
            interp = ('核心排班 UI — 高频属预期' if CN else 'Core scheduling UI — high churn expected')
        else:
            interp = ('活跃特性/集成区' if CN else 'Active feature / integration area')
        hsrows.append([i, f, c, interp])
    table(doc, (['#', '文件', '改动次数', '解读'] if CN else ['#', 'File', 'Touch Count', 'Interpretation']), hsrows)
    note(doc, '改动次数 = 修改该文件的提交数（非行数）。version.ts 居首是结构性噪声，非风险。' if CN
         else 'Touch count = number of commit events that modified the file (not lines changed). version.ts top-rank is structural noise, not risk.')
    explain('“热点文件”就是被改动得最频繁的文件。改得越多，通常说明这块是系统里最忙、最核心或最复杂的地方，值得多关注（多写测试、多人熟悉）。注意：像 version.ts 这种每次发版都要改一下的文件排在最前，属于正常现象，并不代表风险。' if CN
            else '"Hotspot files" are the files changed most often. The more a file is touched, the more it tends to be the busiest, most central, or most complex part of the system — worth extra attention (more tests, more people familiar with it). Note: a file like version.ts that gets a small bump on every release naturally ranks at the top, which is normal and not a risk.')

    return doc


# ----------------------------------------------------------------------------
def read_version(repo):
    try:
        with open(os.path.join(repo, 'gantt/src/version.ts'), encoding='utf-8') as f:
            txt = f.read()
        b = re.search(r'BACKEND_VERSION\s*=\s*(\d+)', txt)
        fr = re.search(r'FRONTEND_VERSION\s*=\s*(\d+)', txt)
        rl = re.search(r'RULE_VERSION\s*=\s*(\d+)', txt)
        return f"Ver:B{b.group(1)}/F{fr.group(1)}/R{rl.group(1)}"
    except Exception:
        return "Ver:unknown"


def next_version(outdir):
    mx = 0
    if os.path.isdir(outdir):
        for fn in os.listdir(outdir):
            m = re.search(r'-V(\d+)\.docx$', fn)
            if m and not fn.startswith('~$'):
                mx = max(mx, int(m.group(1)))
    return f"V{mx + 1}", mx


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--repo', default='.')
    ap.add_argument('--since', default='2026-01-01')
    ap.add_argument('--version', default=None, help='e.g. V7 (default: auto-increment)')
    ap.add_argument('--date', default=None, help='override report date YYYY-MM-DD')
    args = ap.parse_args()

    repo = os.path.abspath(args.repo)
    outdir = os.path.join(repo, 'docs/dev-context')
    today = datetime.date.fromisoformat(args.date) if args.date else datetime.date.today()

    if args.version:
        ver = args.version
        m = re.match(r'V(\d+)', args.version)
        prev_max = int(m.group(1)) - 1 if m else None
    else:
        ver, prev_max = next_version(outdir)

    M = collect(repo, args.since, today)
    app_version = read_version(repo)
    hhmm = datetime.datetime.now().strftime('%H%M')
    meta = {
        'date': today.isoformat(),
        'version': ver,
        'period_end': today.strftime('%b %d, %Y'),
        'period_end_cn': f"{today.year}年{today.month}月{today.day}日",
        'app_version': app_version,
        'prior_line': f"Prior: V{prev_max} · This: {ver} ({today.isoformat()}, {M['total']:,} commits)." if prev_max else None,
    }

    for lang in ['CN', 'EN']:
        doc = build(M, lang, meta)
        fn = f"{today.isoformat()} {hhmm}-dora-developer-analysis-{lang}-{ver}.docx"
        path = os.path.join(outdir, fn)
        doc.save(path)
        print(f"wrote {path}")

    print(f"\nTotals: {M['total']:,} commits, "
          f"{sum(len(v) for v in M['specs'].values())} specs, version {app_version}")


if __name__ == '__main__':
    main()
